"""Test the docx requirements.
"""

# pyTRLCConverter - A tool to convert TRLC files to specific formats.
# Copyright (c) 2024 - 2026 NewTec GmbH
#
# This file is part of pyTRLCConverter program.
#
# The pyTRLCConverter program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License as published by the Free Software Foundation,
# either version 3 of the License, or (at your option) any later version.
#
# The pyTRLCConverter program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
# FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with pyTRLCConverter.
# If not, see <https://www.gnu.org/licenses/>.

# Imports **********************************************************************

from unittest.mock import patch

import docx

from pyTRLCConverter.__main__ import main
from pyTRLCConverter.docx_converter import DocxConverter

# Variables ********************************************************************

# Classes **********************************************************************

# Functions ********************************************************************

def test_tc_docx(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx
    """
    The software shall support conversion of TRLC source files into docx.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx")

    # Mock program arguments to simulate running the script with the built-in docx converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_no_section.trlc",
        "--out", str(tmp_path),
        "docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check that the requirement is present.
    assert created_docx.paragraphs[0].text == "req_id_1 (Requirement)"

    # Check the created table for the requirement.
    assert len(created_docx.tables) == 1
    assert created_docx.tables[0].cell(0, 0).text == "Element"
    assert created_docx.tables[0].cell(0, 1).text == "Value"
    assert created_docx.tables[0].cell(1, 0).text == "description"
    assert created_docx.tables[0].cell(1, 1).text == "Test description"

def test_tc_docx_multiple(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_multiple
    """
    The software shall support conversion of multiple TRLC records containing references to docx format.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_multiple")

    # Mock program arguments to simulate running the script with the built-in docx converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/multi_req_with_link.trlc",
        "--out", str(tmp_path),
        "docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check the link text in the created tables for the requirements.
    assert len(created_docx.tables) == 2
    assert created_docx.tables[0].cell(2, 0).text == "link"
    assert created_docx.tables[0].cell(2, 1).text == "Requirements.req_id_6"
    assert created_docx.tables[1].cell(2, 0).text == "link"
    assert created_docx.tables[1].cell(2, 1).text == "Requirements.req_id_5"

    # Check that the hyperlink anchors were correctly set.
    # Cells paragraph needs to hold a hyperlink with correct anchor.
    assert created_docx.tables[0].cell(2, 1).paragraphs[0]._p.hyperlink_lst[0].anchor == "req_id_6" # pylint: disable=protected-access
    assert created_docx.tables[1].cell(2, 1).paragraphs[0]._p.hyperlink_lst[0].anchor == "req_id_5" # pylint: disable=protected-access

def test_tc_docx_section(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_section
    """
    The test case checks if a TRLC section is correctly converted into docx.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_section")

    # Mock program arguments to simulate running the script with the built-in docx converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_with_section.trlc",
        "--out", str(tmp_path),
        "docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check that the section is present
    assert created_docx.paragraphs[0].text == "Test section"

    # Check that the requirement is present.
    assert created_docx.paragraphs[1].text == "req_id_2 (Requirement)"

def test_tc_docx_file(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_file
    """
    The test case checks whether multiple TRLC files are combined into one docx output document.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_file")

    # Mock program arguments to simulate running the script with the built-in docx converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_no_section.trlc",
        "--source", "./tests/utils/single_req_with_section.trlc",
        "--out", str(tmp_path),
        "docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check that the first requirement is present.
    assert created_docx.paragraphs[0].text == "req_id_1 (Requirement)"
    # Paragraph 1 is the "from file" paragraph.
    # Check the second requirement is present.
    assert created_docx.paragraphs[2].text == "Test section"
    assert created_docx.paragraphs[3].text == "req_id_2 (Requirement)"

def test_tc_docx_template(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_template
    """
    The test case checks whether docx templated are correctly applied.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_template")

    # Mock program arguments to simulate running the script with the built-in docx converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_no_section.trlc",
        "--out", str(tmp_path),
        "docx",
        "--template", "./tests/utils/template.docx",
    ])

    # Expect the program to run without any exceptions.
    main()

    # Capture stdout and stderr.
    captured = capsys.readouterr()
    # Check that no errors were reported.
    assert captured.err == ""

    # Check that the output file was created.
    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # Check that the template was used.
    assert created_docx.paragraphs[0].text == "Template text."

    # Check that the requirement is present.
    assert created_docx.paragraphs[1].text == "req_id_1 (Requirement)"

# pylint: disable-next=too-many-statements, too-many-locals, too-many-branches
def test_tc_docx_render_md(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_render_md
    """
    The test case checks whether strings in CommonMark Markdown are correctly converted to docx.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_render_md")

    # Mock program arguments to simulate running the script with the built-in docx converter.
    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_description_md.trlc",
        "--out", str(tmp_path),
        "--renderCfg", "./tests/utils/renderCfgDocxMd.json",
        "docx",
        "--template", "./tests/utils/template.docx",
    ])

    main()

    captured = capsys.readouterr()
    assert captured.err == ""

    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # The rendered CommonMark content lives in the record attribute table, cell(1,1).
    record_table = created_docx.tables[0]
    description_cell = record_table.rows[1].cells[1]
    cell_texts = [p.text for p in description_cell.paragraphs]

    # Headings rendered as paragraphs inside the cell.
    assert "Heading 1" in cell_texts
    assert "Heading 2" in cell_texts

    # Bullet and numbered list items rendered.
    assert "Bullet point 1" in cell_texts
    assert "Bullet point 2" in cell_texts
    assert "Numbered point 1" in cell_texts
    assert "Numbered point 2" in cell_texts

    # Bold text rendered with bold run property.
    bold_run = None
    for paragraph in description_cell.paragraphs:
        for run in paragraph.runs:
            if run.bold is True and run.text == "Bold text":
                bold_run = run
                break
        if bold_run is not None:
            break

    assert bold_run is not None, "Bold run not found in docx output"

    # Italic text rendered with italic run property.
    italic_run = None
    for paragraph in description_cell.paragraphs:
        for run in paragraph.runs:
            if run.italic is True and run.text == "italic text":
                italic_run = run
                break
        if italic_run is not None:
            break

    assert italic_run is not None, "Italic run not found in docx output"

    # Thematic break rendered as a paragraph with a bottom border (w:pBdr).
    thematic_break_found = False
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for paragraph in description_cell.paragraphs:
        if paragraph._p.findall(f'.//{{{ns}}}pBdr'):  # pylint: disable=protected-access
            thematic_break_found = True
            break

    assert thematic_break_found, "Thematic break paragraph not found in docx output"

    # Blockquote rendered with Quote paragraph style.
    quote_paragraph = None
    for paragraph in description_cell.paragraphs:
        if paragraph.style and paragraph.style.name == "Quote" and paragraph.text == "Blockquote example":
            quote_paragraph = paragraph
            break

    assert quote_paragraph is not None, "Quote paragraph not found in docx output"

    # Hyperlink rendered for the link element.
    hyperlink_found = False
    for paragraph in description_cell.paragraphs:
        if paragraph._p.findall(f'.//{{{ns}}}hyperlink'):  # pylint: disable=protected-access
            hyperlink_found = True
            break

    assert hyperlink_found, "Hyperlink not found in docx output"

# pylint: disable-next=too-many-statements, too-many-locals
def test_tc_docx_render_gfm(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_render_gfm
    """
    The test case checks whether strings in GitHub Flavored Markdown are correctly converted to docx.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Path): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_render_gfm")

    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_description_gfm.trlc",
        "--out", str(tmp_path),
        "--renderCfg", "./tests/utils/renderCfgDocxGfm.json",
        "docx",
        "--template", "./tests/utils/template.docx",
    ])

    main()

    captured = capsys.readouterr()
    assert captured.err == ""

    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    # The rendered GFM content lives in the record attribute table, cell(1,1) — the description value cell.
    record_table = created_docx.tables[0]
    description_cell = record_table.rows[1].cells[1]
    cell_texts = [p.text for p in description_cell.paragraphs]

    # Headings rendered as paragraphs inside the cell.
    assert "Heading 1" in cell_texts
    assert "Heading 2" in cell_texts

    # Bullet and numbered list items rendered.
    assert "Bullet point 1" in cell_texts
    assert "Bullet point 2" in cell_texts
    assert "Numbered point 1" in cell_texts
    assert "Numbered point 2" in cell_texts

    # Bold text rendered with bold run property.
    bold_run = None
    for paragraph in description_cell.paragraphs:
        for run in paragraph.runs:
            if run.bold is True and run.text == "Bold text":
                bold_run = run
                break
        if bold_run is not None:
            break

    assert bold_run is not None, "Bold run not found in docx output"

    # Task list items rendered as list paragraphs with checkbox prefix.
    assert any("☑" in t for t in cell_texts), "Checked task item not found in docx output"
    assert any("☐" in t for t in cell_texts), "Unchecked task item not found in docx output"

    # Thematic break rendered as a paragraph with a bottom border (w:pBdr).
    thematic_break_found = False
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for paragraph in description_cell.paragraphs:
        if paragraph._p.findall(f'.//{{{ns}}}pBdr'):  # pylint: disable=protected-access
            thematic_break_found = True
            break

    assert thematic_break_found, "Thematic break paragraph not found in docx output"

    # Bare URL rendered as a hyperlink.
    bare_url_found = False
    for paragraph in description_cell.paragraphs:
        if paragraph._p.findall(f'.//{{{ns}}}hyperlink'):  # pylint: disable=protected-access
            bare_url_found = True
            break

    assert bare_url_found, "Bare URL hyperlink not found in docx output"

    # Strikethrough rendered with strike font property.
    strike_run = None
    for paragraph in description_cell.paragraphs:
        for run in paragraph.runs:
            if run.font.strike is True and "strikethrough" in run.text:
                strike_run = run
                break
        if strike_run is not None:
            break

    assert strike_run is not None, "Strikethrough run not found in docx output"

    # PlantUML block: if PlantUML is not installed, an error string paragraph is rendered
    # as fallback. Either a picture run or the error string must be present.
    plantuml_rendered = any(
        any(run.element.findall(f'.//{{{ns}}}drawing') for run in p.runs)
        or "[PlantUML error:" in p.text
        for p in description_cell.paragraphs
    )
    assert plantuml_rendered, "PlantUML block not rendered (no image and no error string)"


def test_tc_docx_render_plantuml(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_render_plantuml
    """A plantuml fenced code block in a Markdown attribute shall be rendered as an embedded
    PNG image in the docx output.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Any): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_render_plantuml")

    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_description_md.trlc",
        "--out", str(tmp_path),
        "--renderCfg", "./tests/utils/renderCfgDocxMd.json",
        "docx",
        "--template", "./tests/utils/template.docx",
    ])

    # Provide a minimal valid 1x1 white PNG so python-docx can embed it.
    png_1x1 = (
        b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01'
        b'\x08\x02\x00\x00\x00\x90wS\xde\x00\x00\x00\x0cIDATx\x9cc\xf8\xff\xff?'
        b'\x00\x05\xfe\x02\xfe\r\xefF\xb8\x00\x00\x00\x00IEND\xaeB`\x82'
    )

    def _fake_generate(_fmt, src, _dst):
        png_path = src.replace(".puml", ".png")
        with open(png_path, "wb") as f:
            f.write(png_1x1)

    with patch("pyTRLCConverter.marko.md2docx_renderer.PlantUML.generate",
               side_effect=_fake_generate):
        main()

    captured = capsys.readouterr()
    assert captured.err == ""

    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))
    ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    description_cell = None
    for table in created_docx.tables:
        for row in table.rows:
            if row.cells[0].text == "description":
                description_cell = row.cells[1]
                break
        if description_cell is not None:
            break

    assert description_cell is not None, "description cell not found"

    image_found = any(
        any(run.element.findall(f'.//{{{ns}}}drawing') for run in p.runs)
        for p in description_cell.paragraphs
    )
    assert image_found, "No embedded image found in docx output for plantuml block"


def test_tc_docx_render_plantuml_error(record_property, capsys, monkeypatch, tmp_path):
    # lobster-trace: SwTests.tc_docx_render_plantuml
    """If PlantUML is not available, a plantuml fenced code block shall be replaced by a
    [PlantUML error: ...] paragraph instead of failing the conversion.

    Args:
        record_property (Any): Used to inject the test case reference into the test results.
        capsys (Any): Used to capture stdout and stderr.
        monkeypatch (Any): Used to mock program arguments.
        tmp_path (Any): Used to create a temporary output directory.
    """
    record_property("lobster-trace", "SwTests.tc_docx_render_plantuml")

    monkeypatch.delenv("PLANTUML", raising=False)

    monkeypatch.setattr("sys.argv", [
        "pyTRLCConverter",
        "--source", "./tests/utils/req.rsl",
        "--source", "./tests/utils/single_req_description_md.trlc",
        "--out", str(tmp_path),
        "--renderCfg", "./tests/utils/renderCfgDocxMd.json",
        "docx",
        "--template", "./tests/utils/template.docx",
    ])

    main()

    captured = capsys.readouterr()
    assert captured.err == ""

    created_docx = docx.Document(docx=str(tmp_path / DocxConverter.OUTPUT_FILE_NAME_DEFAULT))

    description_cell = None
    for table in created_docx.tables:
        for row in table.rows:
            if row.cells[0].text == "description":
                description_cell = row.cells[1]
                break
        if description_cell is not None:
            break

    assert description_cell is not None, "description cell not found"

    error_found = any("[PlantUML error:" in p.text for p in description_cell.paragraphs)
    assert error_found, "No [PlantUML error:] paragraph found in docx output"

# Main *************************************************************************
