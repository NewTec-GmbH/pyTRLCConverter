"""DOCX renderer for Marko.
    It is used to convert GitHub Flavored Markdown AST to DOCX format.

    Author: Stefan Vogel (stefan.vogel@newtec.de)
"""

# pyTRLCConverter - A tool to convert TRLC files to specific formats.
# Copyright (c) 2024 - 2026 NewTec GmbH
#
# This file is part of pyTRLCConverter program.
#
# The pyTRLCConverter program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License as published by the Free Software Foundation,
# either version 3 of the License, or (at your option) any later version.
#
# The pyTRLCConverter program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
# FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with pyTRLCConverter.
# If not, see <https://www.gnu.org/licenses/>.

# Imports **********************************************************************

from __future__ import annotations
from typing import Any, Optional
from docx.blkcntnr import BlockItemContainer
from pyTRLCConverter.marko.md2docx_renderer import Md2DocxRenderer, Singleton

# Variables ********************************************************************

# Classes **********************************************************************

class Gfm2DocxRenderer(Md2DocxRenderer, metaclass=Singleton):
    # lobster-trace: SwRequirements.sw_req_docx_render_gfm
    """Renderer for DOCX output.
    It is used to convert GitHub Flavored Markdown to DOCX format.
    """

    # DOCX block item container to add content to.
    block_item_container: Optional[BlockItemContainer] = None

    def __init__(self) -> None:
        """Initialize the renderer."""
        self._checkbox_prefix: Optional[str] = None
        super().__init__()
        self.reset()

    def reset(self) -> None:
        """Resets the renderer state before a new convert() call."""
        super().reset()
        self._checkbox_prefix = None

    def render_paragraph(self, element: Any) -> None:
        """Renders a paragraph element, prepending a checkbox prefix if set by render_list_item.

        Args:
            element (Any): The paragraph element to render.
        """
        assert self.block_item_container is not None

        # If a checkbox prefix is pending, create the paragraph manually so we can add the
        # prefix run before rendering inline children — super() would create the paragraph
        # and immediately render children leaving no hook to insert the prefix first.
        if self._checkbox_prefix is not None:
            from docx.oxml import OxmlElement  # pylint: disable=import-outside-toplevel
            from docx.oxml.ns import qn as qname  # pylint: disable=import-outside-toplevel

            style = self._list_style[-1] if self._list_style else "List Bullet"
            self._current_paragraph = self.block_item_container.add_paragraph(style=style)  # type: ignore

            r = OxmlElement('w:r')
            t = OxmlElement('w:t')
            t.text = self._checkbox_prefix
            t.set(qname('xml:space'), 'preserve')
            r.append(t)
            self._current_paragraph._p.append(r)  # pylint: disable=protected-access

            self._checkbox_prefix = None
            self.render_children(element)
            self._current_paragraph = None
        else:
            super().render_paragraph(element)

    def render_list_item(self, element: Any) -> None:
        """Renders a list item element, prepending a checkbox prefix for GFM task list items.

        Args:
            element (Any): The list item element to render.
        """
        # GFM task-list items expose a "checked" flag on the first child Paragraph node.
        # Store the prefix as state so render_paragraph can prepend it as a run, keeping
        # the checkbox and text in the same paragraph rather than creating a separate one.
        if element.children:
            first_child = element.children[0]
            checked = getattr(first_child, "checked", None)
            if checked is True:
                self._checkbox_prefix = "☑ "
            elif checked is False:
                self._checkbox_prefix = "☐ "

        self.render_children(element)

    def render_table(self, element: Any) -> None:
        """Renders a GFM table as a python-docx table.

        Args:
            element (Any): The GFM table element.
        """
        assert self.block_item_container is not None

        if element.children:
            header_row = element.children[0]
            body_rows = element.children[1:]

            num_cols = len(header_row.children)
            num_rows = 1 + len(body_rows)

            table = self.block_item_container.add_table(rows=num_rows, cols=num_cols) # type: ignore
            table.style = "Table Grid"

            # Header row
            for col_idx, cell in enumerate(header_row.children):
                cell_text = self._collect_text(cell)
                table.rows[0].cells[col_idx].text = cell_text
                for run in table.rows[0].cells[col_idx].paragraphs[0].runs:
                    run.bold = True

            # Body rows
            for row_idx, row in enumerate(body_rows):
                for col_idx, cell in enumerate(row.children):
                    cell_text = self._collect_text(cell)
                    table.rows[row_idx + 1].cells[col_idx].text = cell_text

            # python-docx inserts an empty paragraph after every table added to a cell
            # as an OOXML structural requirement. Remove it to avoid a spurious blank line
            # between the table and the next content element.
            tbl = table._tbl  # pylint: disable=protected-access
            tbl_following = tbl.getnext()
            if tbl_following is not None and tbl_following.tag.endswith("}p"):
                tbl_following.getparent().remove(tbl_following)

    def render_table_row(self, element: Any) -> None:
        """No-op: table rows are handled inside render_table.

        Args:
            element (Any): The table row element.
        """

    def render_table_cell(self, element: Any) -> None:
        """No-op: table cells are handled inside render_table.

        Args:
            element (Any): The table cell element.
        """

    def render_strikethrough(self, element: Any) -> None:
        """Renders a GFM strikethrough element by applying the strike run property.

        Args:
            element (Any): The strikethrough element.
        """
        assert self._current_paragraph is not None

        run = self._current_paragraph.add_run(self._collect_text(element))
        run.font.strike = True

    def render_url(self, element: Any) -> None:
        """Renders a GFM bare URL by delegating to render_auto_link.

        Args:
            element (Any): The URL element.
        """
        self.render_auto_link(element)

    @staticmethod
    def _collect_text(element: Any) -> str:
        """Recursively collects plain text from an element's children.

        Args:
            element (Any): The element to collect text from.

        Returns:
            str: The concatenated plain text of the element.
        """
        if isinstance(element.children, str):
            result = element.children
        else:
            parts = []
            for child in element.children:
                parts.append(Gfm2DocxRenderer._collect_text(child))
            result = "".join(parts)

        return result

# Functions ********************************************************************

# Main *************************************************************************
