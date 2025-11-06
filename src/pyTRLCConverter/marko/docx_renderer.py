"""Docx Renderer for Marko.

    Author: Andreas Merkle (andreas.merkle@newtec.de)
"""

# pyTRLCConverter - A tool to convert TRLC files to specific formats.
# Copyright (c) 2024 - 2025 NewTec GmbH
#
# This file is part of pyTRLCConverter program.
#
# The pyTRLCConverter program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License as published by the Free Software Foundation,
# either version 3 of the License, or (at your option) any later version.
#
# The pyTRLCConverter program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
# FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with pyTRLCConverter.
# If not, see <https://www.gnu.org/licenses/>.

# Imports **********************************************************************

from __future__ import annotations
from typing import TYPE_CHECKING, Any, cast, Optional
from marko import Renderer
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.blkcntnr import BlockItemContainer

if TYPE_CHECKING:
    from . import block, inline

# Variables ********************************************************************

# Classes **********************************************************************

class Singleton(type):
    # lobster-trace: SwRequirements.sw_req_docx_render_md
    """Singleton metaclass to ensure only one instance of a class exists."""

    _instances = {}

    def __call__(cls, *args, **kwargs):
        """
        Returns the singleton instance of the class.

        Returns:
            instance (cls): The singleton instance of the class.
        """

        if cls not in cls._instances:
            cls._instances[cls] = super(Singleton, cls).__call__(*args, **kwargs)

        return cls._instances[cls]

# pylint: disable-next=too-many-public-methods, too-many-instance-attributes
class DocxRenderer(Renderer, metaclass=Singleton):
    # lobster-trace: SwRequirements.sw_req_docx_render_md
    """Renderer for docx output."""

    # Docx block item container to add content to.
    block_item_container: Optional[BlockItemContainer] = None

    def __init__(self) -> None:
        """Initialize docx renderer."""
        super().__init__()
        self._list_indent_level = 0
        self._is_italic = False
        self._is_bold = False
        self._is_underline = False
        self._is_heading = False
        self._heading_level = 0
        self._is_list_item = False
        self._list_style = []
        self._is_quote = False

    def render_children(self, element: Any) -> None:
        """
        Recursively renders child elements of a given element to
        a docx document.

        Args:
            element (Element): The parent element whose children are to be rendered.
        """
        for child in element.children:
            self.render(child)

    def render_paragraph(self, element: block.Paragraph) -> None:
        """
        Renders a paragraph element.

        Args:
            element (block.Paragraph): The paragraph element to render.
        """
        self.render_children(element)

    def render_list(self, element: block.List) -> None:
        """
        Renders a list (ordered or unordered) element.

        Args:
            element (block.List): The list element to render.
        """
        assert self.block_item_container is not None

        self._is_list_item = True
        self._list_indent_level += 1

        style = "List Number" if element.ordered else "List Bullet"

        if self._list_indent_level > 1:
            style += f" {self._list_indent_level}"

        self._list_style.append(style)

        for child in element.children:
            self.render_children(child)

        self._list_style.pop()

        self._list_indent_level -= 1

        if self._list_indent_level == 0:
            self._is_list_item = False

    def render_quote(self, element: block.Quote) -> None:
        """
        Renders a blockquote element.

        Args:
            element (block.Quote): The blockquote element to render.
        """
        self._is_quote = True
        self.render_children(element)
        self._is_quote = False

    def render_fenced_code(self, element: block.FencedCode) -> None:
        """
        Renders a fenced code block element.

        Args:
            element (block.FencedCode): The fenced code block element to render.
        """
        assert self.block_item_container is not None

        paragraph = self.block_item_container.add_paragraph()
        run = paragraph.add_run(element.children[0].children)
        run.font.name = "Consolas"

    def render_code_block(self, element: block.CodeBlock) -> None:
        """
        Renders a code block element.

        Args:
            element (block.CodeBlock): The code block element to render.
        """
        self.render_fenced_code(cast("block.FencedCode", element))

    def render_html_block(self, element: block.HTMLBlock) -> None:
        """
        Renders a raw HTML block element.

        Args:
            element (block.HTMLBlock): The HTML block element to render.
        """
        self.render_fenced_code(cast("block.FencedCode", element))

    # pylint: disable-next=unused-argument
    def render_thematic_break(self, element: block.ThematicBreak) -> None:
        """
        Renders a thematic break (horizontal rule) element.

        Args:
            element (block.ThematicBreak): The thematic break element to render.
        """
        assert self.block_item_container is not None

        # Add a horizontal rule by inserting a paragraph with a bottom border
        para = self.block_item_container.add_paragraph()
        p = para._element # pylint: disable=protected-access

        paragraph_properties = p.get_or_add_pPr()
        p_bdr = OxmlElement('w:pBdr')
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '6')
        bottom.set(qn('w:space'), '1')
        bottom.set(qn('w:color'), 'auto')
        p_bdr.append(bottom)
        paragraph_properties.append(p_bdr)

    def render_heading(self, element: block.Heading) -> None:
        """
        Renders a heading element.

        Args:
            element (block.Heading): The heading element to render.
        """
        assert self.block_item_container is not None

        self._is_heading = True
        self._heading_level = min(max(element.level, 1), 9)  # docx supports levels 1-9

        self.render_children(element)

        self._is_heading = False
        self._heading_level = 0

    def render_setext_heading(self, element: block.SetextHeading) -> None:
        """
        Renders a setext heading element.

        Args:
            element (block.SetextHeading): The setext heading element to render.
        """
        self.render_heading(cast("block.Heading", element))

    # pylint: disable-next=unused-argument
    def render_blank_line(self, element: block.BlankLine) -> None:
        """
        Renders a blank line element.

        Args:
            element (block.BlankLine): The blank line element to render.
        """
        assert self.block_item_container is not None

        paragraph = self.block_item_container.add_paragraph()
        run = paragraph.add_run("")
        run.bold = self._is_bold
        run.italic = self._is_italic
        run.underline = self._is_underline

        if self._is_heading is True:
            paragraph.style = f"Heading {self._heading_level}"
        elif self._is_list_item:
            paragraph.style = self._list_style[-1]

    # pylint: disable-next=unused-argument
    def render_link_ref_def(self, element: block.LinkRefDef) -> None:
        """
        Renders a link reference definition element.

        Args:
            element (block.LinkRefDef): The link reference definition element to render.
        """
        # reStructuredText uses reference links differently than Markdown.
        # It shall not be rendered in the document.

    def render_emphasis(self, element: inline.Emphasis) -> None:
        """
        Renders an emphasis (italic) element.

        Args:
            element (inline.Emphasis): The emphasis element to render.
        """
        assert self.block_item_container is not None

        self._is_italic = True

        self.render_children(element)

        self._is_italic = False

    def render_strong_emphasis(self, element: inline.StrongEmphasis) -> None:
        """
        Renders a strong emphasis (bold) element.

        Args:
            element (inline.StrongEmphasis): The strong emphasis element to render.
        """
        assert self.block_item_container is not None

        self._is_bold = True

        self.render_children(element)

        self._is_bold = False

    def render_inline_html(self, element: inline.InlineHTML) -> None:
        """
        Renders an inline HTML element.

        Args:
            element (inline.InlineHTML): The inline HTML element to render.
        """
        assert self.block_item_container is not None

        paragraph = self.block_item_container.add_paragraph()
        run = paragraph.add_run(element.children)
        run.font.name = "Consolas"

    def render_plain_text(self, element: Any) -> None:
        """
        Renders plain text or any element with string children.

        Args:
            element (Any): The element to render.
        """
        assert self.block_item_container is not None

        if isinstance(element.children, str):
            paragraph = self.block_item_container.add_paragraph()
            run = paragraph.add_run(element.children)
            run.bold = self._is_bold
            run.italic = self._is_italic
            run.underline = self._is_underline

            if self._is_heading is True:
                paragraph.style = f"Heading {self._heading_level}"
            elif self._is_list_item:
                paragraph.style = self._list_style[-1]
        else:
            self.render_children(element)

    def render_link(self, element: inline.Link) -> None:
        """
        Renders a link element.

        Args:
            element (inline.Link): The link element to render.
        """
        # Link handling in docx is non-trivial; for simplicity, render link text only.
        self.render_children(element)

    def render_auto_link(self, element: inline.AutoLink) -> None:
        """
        Renders an auto link element.

        Args:
            element (inline.AutoLink): The auto link element to render.
        """
        self.render_link(cast("inline.Link", element))

    def render_image(self, element: inline.Image) -> None:
        """
        Renders an image element.

        Args:
            element (inline.Image): The image element to render.

        Returns:
            DocxDocument: The rendered image as a docx document.
        """
        # Image handling in docx is non-trivial; for simplicity just render title and URL.
        if element.title:
            assert self.block_item_container is not None
            self.block_item_container.add_paragraph(text=element.title)
            self.block_item_container.add_paragraph(text=f" ({element.dest})")

    def render_literal(self, element: inline.Literal) -> None:
        """
        Renders a literal (inline code) element.

        Args:
            element (inline.Literal): The literal element to render.
        """
        self.render_raw_text(cast("inline.RawText", element))

    def render_raw_text(self, element: inline.RawText) -> None:
        """
        Renders a raw text element.

        Args:
            element (inline.RawText): The raw text element to render.

        Returns:
            DocxDocument: The rendered raw text as a docx document.
        """
        assert self.block_item_container is not None

        paragraph = self.block_item_container.add_paragraph()
        run = paragraph.add_run(element.children)
        run.bold = self._is_bold
        run.italic = self._is_italic
        run.underline = self._is_underline

        if self._is_heading is True:
            paragraph.style = f"Heading {self._heading_level}"
        elif self._is_list_item is True:
            paragraph.style = self._list_style[-1]
        elif self._is_quote is True:
            paragraph.style = "Quote"

    # pylint: disable-next=unused-argument
    def render_line_break(self, element: inline.LineBreak) -> None:
        """
        Renders a line break element.

        Args:
            element (inline.LineBreak): The line break element to render.

        Returns:
            DocxDocument: The rendered line break as a docx document.
        """
        assert self.block_item_container is not None

        paragraph = self.block_item_container.add_paragraph()
        run = paragraph.add_run()
        run.add_break()

    def render_code_span(self, element: inline.CodeSpan) -> None:
        """
        Renders a code span (inline code) element.

        Args:
            element (inline.CodeSpan): The code span element to render.

        Returns:
            DocxDocument: The rendered code span as a docx document.
        """
        assert self.block_item_container is not None

        paragraph = self.block_item_container.add_paragraph()
        run = paragraph.add_run(cast(str, element.children))
        run.font.name = "Consolas"

# Functions ********************************************************************

# Main *************************************************************************
