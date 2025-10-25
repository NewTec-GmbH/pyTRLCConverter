"""Converter to Word docx format.

    Author: Norbert Schulz (norbert.schulz@newtec.de)
"""

# pyTRLCConverter - A tool to convert TRLC files to specific formats.
# Copyright (c) 2024 - 2025 NewTec GmbH
#
# This file is part of pyTRLCConverter program.
#
# The pyTRLCConverter program is free software: you can redistribute it and/or modify it under
# the terms of the GNU General Public License as published by the Free Software Foundation,
# either version 3 of the License, or (at your option) any later version.
#
# The pyTRLCConverter program is distributed in the hope that it will be useful, but
# WITHOUT ANY WARRANTY; without even the implied warranty of MERCHANTABILITY or
# FITNESS FOR A PARTICULAR PURPOSE. See the GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License along with pyTRLCConverter.
# If not, see <https://www.gnu.org/licenses/>.

# Imports **********************************************************************
import os
from typing import Optional, Any
import docx
from docx.blkcntnr import BlockItemContainer
from docx.text.paragraph import Paragraph
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from marko import Markdown
from trlc.ast import Implicit_Null, Record_Object, Record_Reference, String_Literal, Array_Aggregate, Expression
from pyTRLCConverter.base_converter import BaseConverter
from pyTRLCConverter.marko.docx_renderer import DocxRenderer
from pyTRLCConverter.ret import Ret
from pyTRLCConverter.trlc_helper import TrlcAstWalker
from pyTRLCConverter.logger import log_verbose

# Variables ********************************************************************

# Classes **********************************************************************

class DocxConverter(BaseConverter):
    """
    Converter to docx format.

    The following Word docx objects are used:
    - Document: Represents the entire Word document. You can create a new document or load an existing one.
    - Paragraph: A block of text in the document. It has its own formatting properties.
    - Run: A contiguous run of text with the same formatting. You can change the formatting of a run
            independently within a paragraph.
    - Table: A two-dimensional structure for presenting data in rows and columns.
    """

    OUTPUT_FILE_NAME_DEFAULT = "output.docx"

    def __init__(self, args: Any) -> None:
        # lobster-trace: SwRequirements.sw_req_no_prj_spec
        # lobster-trace: SwRequirements.sw_req_docx
        # lobster-trace: SwRequirements.sw_req_docx_template
        """
        Initialize the docx converter.

        Args:
            args (Any): The parsed program arguments.
        """
        super().__init__(args)

        if args.template is not None:
            log_verbose(f"Loading template file {args.template}.")

        self._docx = docx.Document(docx=args.template)

        # Ensure default table style is present in the document.
        if not 'Table Grid' in self._docx.styles:
            self._docx.styles.add_style('Table Grid', docx.enum.style.WD_STYLE_TYPE.TABLE, builtin=True)

        # The AST walker meta data for processing the record object fields.
        # This will hold the information about the current package, type and attribute being processed.
        self._ast_meta_data = None

        # Current list item indentation level.
        self._list_item_indent_level = 0

        # Docx block item container to add content to during conversion and markdown rendering.
        self._block_item_container: Optional[BlockItemContainer] = None

    @staticmethod
    def get_subcommand() -> str:
        # lobster-trace: SwRequirements.sw_req_docx
        """ Return subcommand token for this converter.

        Returns:
            Ret: Status
        """
        return "docx"

    @staticmethod
    def get_description() -> str:
        # lobster-trace: SwRequirements.sw_req_docx
        """ Return converter description.
 
        Returns:
            Ret: Status
        """
        return "Convert into docx format."

    @classmethod
    def register(cls, args_parser: Any) -> None:
        # lobster-trace: SwRequirements.sw_req_docx
        """Register converter specific argument parser.

        Args:
            args_parser (Any): Argument parser
        """
        super().register(args_parser)

        BaseConverter._parser.add_argument(
            "-t",
            "--template",
            type=str,
            default=None,
            required=False,
            help="Load the given docx file as a template to append to."
        )
        BaseConverter._parser.add_argument(
            "-n",
            "--name",
            type=str,
            default=DocxConverter.OUTPUT_FILE_NAME_DEFAULT,
            required=False,
            help="Name of the generated output file inside the output folder " \
                f"(default = {DocxConverter.OUTPUT_FILE_NAME_DEFAULT})."
        )

    def convert_section(self, section: str, level: int) -> Ret:
        # lobster-trace: SwRequirements.sw_req_docx_section
        """Process the given section item.

        Args:
            section (str): The section name
            level (int): The section indentation level

        Returns:
            Ret: Status
        """
        self._docx.add_heading(section, level)

        return Ret.OK

    def convert_record_object_generic(self, record: Record_Object, level: int, translation: Optional[dict]) -> Ret:
        # lobster-trace: SwRequirements.sw_req_docx_record
        """
        Process the given record object in a generic way.

        The handler is called by the base converter if no specific handler is
        defined for the record type.

        Args:
            record (Record_Object): The record object.
            level (int): The record level.
            translation (Optional[dict]): Translation dictionary for the record object.
                                            If None, no translation is applied.

        
        Returns:
            Ret: Status
        """
        return self._convert_record_object(record, level, translation)

    def finish(self) -> Ret:
        # lobster-trace: SwRequirements.sw_req_docx_file
        """Finish the conversion.

        Returns:
            Ret: Status
        """
        result = Ret.ERROR

        if self._docx is not None:
            output_file_name = self._args.name
            if 0 < len(self._args.out):
                output_file_name = os.path.join(self._args.out, self._args.name)

            log_verbose(f"Writing docx {output_file_name}.")
            self._docx.save(output_file_name)
            self._docx = None
            result = Ret.OK

        return result

    def _on_implict_null(self, _: Implicit_Null) -> None:
        # lobster-trace: SwRequirements.sw_req_docx_record
        """
        Process the given implicit null value.        
        """
        assert self._block_item_container is not None
        self._block_item_container.add_paragraph(self._empty_attribute_value)

    def _on_record_reference(self, record_reference: Record_Reference) -> None:
        # lobster-trace: SwRequirements.sw_req_docx_record
        """
        Process the given record reference value and return a hyperlink paragraph.

        Args:
            record_reference (Record_Reference): The record reference value.        
        """
        assert self._block_item_container is not None
        paragraph = self._block_item_container.add_paragraph()

        DocxConverter.docx_add_link_to_bookmark(paragraph,
                                                record_reference.target.name,
                                                f"{record_reference.package.name}.{record_reference.target.name}")

    def _on_string_literal(self, string_literal: String_Literal) -> None:
        # lobster-trace: SwRequirements.sw_req_docx_string_format
        """
        Process the given string literal value.

        Args:
            string_literal (String_Literal): The string literal value.
        """
        assert self._block_item_container is not None

        is_handled = False

        if self._ast_meta_data is not None:
            package_name = self._ast_meta_data.get("package_name", "")
            type_name = self._ast_meta_data.get("type_name", "")
            attribute_name = self._ast_meta_data.get("attribute_name", "")

            self._render(package_name, type_name, attribute_name, string_literal.to_string())
            is_handled = True

        if is_handled is False:
            self._block_item_container.add_paragraph(string_literal.to_string())

    # pylint: disable-next=unused-argument
    def _on_array_aggregate_begin(self, array_aggregate: Array_Aggregate) -> None:
        """
        Handle the beginning of a list.

        Args:
            array_aggregate (Array_Aggregate): The AST node.
        """
        self._list_item_indent_level += 1

    # pylint: disable-next=unused-argument
    def _on_list_item(self, expression: Expression, item_result: Any) -> Any:
        # lobster-trace: SwRequirements.sw_req_docx_record
        """
        Handle the list item by adding a bullet point.

        Args:
            expression (Expression): The AST node.
            item_result (Union[list[DocumentObject],DocumentObject]): The result of processing the list item.

        Returns:
            Any: The processed list item.
        """
        assert self._block_item_container is not None

        # Add list item style to last added paragraph.
        last_paragraph = self._block_item_container.paragraphs[-1]

        style = 'List Bullet'

        if 1 < self._list_item_indent_level:
            style += f' {self._list_item_indent_level}'

        last_paragraph.style = style

        return item_result

    # pylint: disable-next=unused-argument
    def _on_array_aggregate_finish(self, array_aggregate: Array_Aggregate) -> None:
        """
        Handle the end of a list.

        Args:
            array_aggregate (Array_Aggregate): The AST node.
        """
        self._list_item_indent_level -= 1

    def _other_dispatcher(self, expression: Expression) -> None:
        """
        Dispatcher for all other expressions.

        Args:
            expression (Expression): The expression to process.
        """
        assert self._block_item_container is not None
        self._block_item_container.add_paragraph(expression.to_string())

    def _get_trlc_ast_walker(self) -> TrlcAstWalker:
        # lobster-trace: SwRequirements.sw_req_docx_record
        # lobster-trace: SwRequirements.sw_req_docx_string_format
        """
        If a record object contains a record reference, the record reference will be converted to
        a hyperlink.
        If a record object contains an array of record references, the array will be converted to
        a list of links.
        Otherwise the record object fields attribute values will be written to the table.

        Returns:
            TrlcAstWalker: The TRLC AST walker.
        """
        trlc_ast_walker = TrlcAstWalker()
        trlc_ast_walker.add_dispatcher(
            Implicit_Null,
            None,
            self._on_implict_null,
            None
        )
        trlc_ast_walker.add_dispatcher(
            Record_Reference,
            None,
            self._on_record_reference,
            None
        )
        trlc_ast_walker.add_dispatcher(
            String_Literal,
            None,
            self._on_string_literal,
            None
        )
        trlc_ast_walker.add_dispatcher(
            Array_Aggregate,
            self._on_array_aggregate_begin,
            None,
            self._on_array_aggregate_finish
        )
        trlc_ast_walker.set_other_dispatcher(self._other_dispatcher)
        trlc_ast_walker.set_list_item_dispatcher(self._on_list_item)

        return trlc_ast_walker

    def _render(self, package_name: str, type_name: str, attribute_name: str, attribute_value: str) -> None:
        # lobster-trace: SwRequirements.sw_req_rst_string_format
        # lobster-trace: SwRequirements.sw_req_rst_string_format_md
        """Render the attribute value depened on its format.

        Args:
            package_name (str): The package name.
            type_name (str): The type name.
            attribute_name (str): The attribute name.
            attribute_value (str): The attribute value.
        """
        assert self._block_item_container is not None

        # If the attribute is marked as markdown format, convert it.
        if self._render_cfg.is_format_md(package_name, type_name, attribute_name) is True:
            DocxRenderer.block_item_container = self._block_item_container
            markdown = Markdown(renderer=DocxRenderer)
            markdown.convert(attribute_value)
        else:
            self._block_item_container.add_paragraph(attribute_value)

    def _convert_record_object(self, record: Record_Object, level: int, translation: Optional[dict]) -> Ret:
        # lobster-trace: SwRequirements.sw_req_docx_record
        """
        Process the given record object.

        Args:
            record (Record_Object): The record object.
            level (int): The record level.
            translation (Optional[dict]): Translation dictionary for the record object.
                                            If None, no translation is applied.

        Returns:
            Ret: Status
        """
        heading = self._docx.add_heading(f"{record.name} ({record.n_typ.name})", level + 1)
        DocxConverter.docx_add_bookmark(heading, record.name)

        table = self._docx.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.autofit = True

        # Set table headers
        header_cells = table.rows[0].cells
        header_cells[0].text = "Element"
        header_cells[1].text = "Value"

        # Walk through the record object fields and write the table rows.
        trlc_ast_walker = self._get_trlc_ast_walker()

        for name, value in record.field.items():
            attribute_name = self._translate_attribute_name(translation, name)

            cells = table.add_row().cells
            cells[0].text = attribute_name

            self._ast_meta_data = {
                "package_name": record.n_package.name,
                "type_name": record.n_typ.name,
                "attribute_name": name
            }
            self._block_item_container = cells[1]
            trlc_ast_walker.walk(value)

            # Remove first empty paragraph added by default to the table cell.
            if 1 < len(cells[1].paragraphs):
                first_paragraph = cells[1].paragraphs[0]

                if first_paragraph.text == "":
                    p_element = first_paragraph._element # pylint: disable=protected-access
                    p_element.getparent().remove(p_element)
                    p_element._p = p_element._element = None # pylint: disable=protected-access

        # Add a paragraph with the record object location
        paragraph = self._docx.add_paragraph()
        paragraph.add_run(f"from {record.location.file_name}:{record.location.line_no}").italic = True

        return Ret.OK

    @staticmethod
    def docx_add_bookmark(paragraph: Paragraph, bookmark_name: str) -> None:
        """
        Adds a bookmark to a paragraph.

        Args:
            paragraph (Paragraph): The paragraph to add the bookmark to.
            bookmark_name (str): The name of the bookmark.
        """
        element = paragraph._p # pylint: disable=protected-access

        # Create a bookmark start element.
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), '0')  # ID must be unique
        bookmark_start.set(qn('w:name'), bookmark_name)

        # Create a bookmark end element.
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), '0')

        # Add the bookmark to the paragraph.
        element.insert(0, bookmark_start)
        element.append(bookmark_end)

    @staticmethod
    def docx_add_link_to_bookmark(paragraph: Paragraph, bookmark_name: str, link_text: str) -> None:
        """
        Add a hyperlink to a bookmark in a paragraph.

        Args:
            paragraph (Paragraph): The paragraph to add the hyperlink to.
            bookmark_name (str): The name of the bookmark.
            link_text (str): The text to display for the hyperlink.
        """
        # Create hyperlink element pointing to the bookmark.
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('w:anchor'), bookmark_name)

        # Create a run and run properties for the hyperlink.
        new_run = OxmlElement('w:r')
        run_properties = OxmlElement('w:rPr')

        # Use the built-in Hyperlink run style so Word will display it correctly (blue/underline).
        r_style = OxmlElement('w:rStyle')
        r_style.set(qn('w:val'), 'Hyperlink')
        run_properties.append(r_style)

        new_run.append(run_properties)

        # Add the text node inside the run (w:t).
        text_element = OxmlElement('w:t')
        text_element.text = link_text
        new_run.append(text_element)

        hyperlink.append(new_run)

        # Append the hyperlink element directly to the paragraph XML so Word renders it.
        paragraph._p.append(hyperlink)  # pylint: disable=protected-access

# Functions ********************************************************************

# Main *************************************************************************
